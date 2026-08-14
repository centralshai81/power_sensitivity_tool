from __future__ import annotations

from pathlib import Path


def export_report_to_docx(markdown_text: str, output_path: str | Path) -> bool:
    """
    将 Markdown 文本粗略导出为 docx。
    依赖 python-docx；若环境无该库，返回 False。
    """
    try:
        from docx import Document
    except Exception:
        return False

    p = Path(output_path)
    p.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip() == "```":
            continue
        else:
            doc.add_paragraph(line)

    doc.save(p)
    return True
